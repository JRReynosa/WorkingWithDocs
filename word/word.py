import docx

doc = docx.Document('demo.docx')
len(doc.paragraphs)
# Output; 7

doc.paragraphs[0].text
# Output; 'Document Title'

doc.paragraphs[1].text
# Output; 'A plain paragraph with some bold and some italic'

len(doc.paragraphs[1].runs)
# Output; 4

doc.paragraphs[1].runs[0].text
# Output; 'A plain paragraph with some '
doc.paragraphs[1].runs[1].text
# Output; 'bold'
doc.paragraphs[1].runs[2].text
# Output; ' and some '
doc.paragraphs[1].runs[3].text
# Output; 'italic'


# Getting full text
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))
# Output:
# Document Title
# A plain paragraph with some bold and some italic
# Heading, level 1
# Intense quote
# first item in unordered list
# first item in ordered list

# Writing Word Docs
doc = docx.Document()
doc.add_paragraph('Hello, world!')
# Output: <docx.text.Paragraph object at 0x0000000003B56F60>

paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')

doc.save('helloworld.docx')